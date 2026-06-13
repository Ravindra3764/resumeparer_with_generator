from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


HEADING_COLOR = RGBColor(0x1F, 0x3A, 0x5F)


def _set_margins(doc, inches=0.7):
    for section in doc.sections:
        section.top_margin = Inches(inches)
        section.bottom_margin = Inches(inches)
        section.left_margin = Inches(inches)
        section.right_margin = Inches(inches)


def _add_bottom_border(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1F3A5F')
    pBdr.append(bottom)
    pPr.append(pBdr)


def _section_heading(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = HEADING_COLOR
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    _add_bottom_border(p)
    return p


def _bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(text)
    p.paragraph_format.space_after = Pt(2)
    for run in p.runs:
        run.font.size = Pt(10.5)
    return p


def generate_resume_docx(data, output_path):
    doc = Document()
    _set_margins(doc)

    # Base font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10.5)

    # Name header
    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_p.add_run(data.get("name", "").strip() or "Your Name")
    name_run.bold = True
    name_run.font.size = Pt(20)
    name_run.font.color.rgb = HEADING_COLOR
    name_p.paragraph_format.space_after = Pt(2)

    # Contact line
    contact = data.get("contact", {})
    contact_parts = [contact.get(k, "") for k in ("email", "phone", "location", "linkedin", "other")]
    contact_parts = [c for c in contact_parts if c]
    if contact_parts:
        contact_p = doc.add_paragraph()
        contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = contact_p.add_run(" | ".join(contact_parts))
        run.font.size = Pt(9.5)
        contact_p.paragraph_format.space_after = Pt(6)

    # Summary
    if data.get("summary"):
        _section_heading(doc, "Professional Summary")
        p = doc.add_paragraph(data["summary"])
        p.paragraph_format.space_after = Pt(4)
        for run in p.runs:
            run.font.size = Pt(10.5)

    # Skills
    if data.get("skills"):
        _section_heading(doc, "Skills")
        p = doc.add_paragraph(" \u2022 ".join(data["skills"]))
        p.paragraph_format.space_after = Pt(4)
        for run in p.runs:
            run.font.size = Pt(10.5)

    # Experience
    if data.get("experience"):
        _section_heading(doc, "Experience")
        for job in data["experience"]:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            left = p.add_run(f"{job.get('title','')} \u2014 {job.get('company','')}")
            left.bold = True
            left.font.size = Pt(11)

            # tab to right-align dates
            tab_stops = p.paragraph_format.tab_stops
            tab_stops.add_tab_stop(Inches(6.6))
            p.add_run("\t")
            date_loc = " | ".join(filter(None, [job.get("location", ""), job.get("dates", "")]))
            date_run = p.add_run(date_loc)
            date_run.italic = True
            date_run.font.size = Pt(9.5)

            for bullet in job.get("bullets", []):
                _bullet(doc, bullet)

    # Projects
    if data.get("projects"):
        _section_heading(doc, "Projects")
        for proj in data["projects"]:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(proj.get("name", ""))
            run.bold = True
            run.font.size = Pt(11)
            if proj.get("description"):
                desc_p = doc.add_paragraph(proj["description"])
                desc_p.paragraph_format.space_after = Pt(2)
                for run in desc_p.runs:
                    run.font.size = Pt(10.5)
            for bullet in proj.get("bullets", []):
                _bullet(doc, bullet)

    # Education
    if data.get("education"):
        _section_heading(doc, "Education")
        for edu in data["education"]:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            left = p.add_run(f"{edu.get('degree','')} \u2014 {edu.get('school','')}")
            left.bold = True
            left.font.size = Pt(10.5)

            tab_stops = p.paragraph_format.tab_stops
            tab_stops.add_tab_stop(Inches(6.6))
            p.add_run("\t")
            date_loc = " | ".join(filter(None, [edu.get("location", ""), edu.get("dates", "")]))
            date_run = p.add_run(date_loc)
            date_run.italic = True
            date_run.font.size = Pt(9.5)

    # Certifications
    if data.get("certifications"):
        _section_heading(doc, "Certifications")
        for cert in data["certifications"]:
            _bullet(doc, cert)

    doc.save(output_path)
    return output_path


def generate_cover_letter_docx(resume_data, body_text, company_name, output_path):
    doc = Document()
    _set_margins(doc, inches=1.0)

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    contact = resume_data.get("contact", {})

    # Sender header
    name_p = doc.add_paragraph()
    run = name_p.add_run(resume_data.get("name", "").strip() or "Your Name")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = HEADING_COLOR

    contact_parts = [contact.get(k, "") for k in ("email", "phone", "location", "linkedin")]
    contact_parts = [c for c in contact_parts if c]
    if contact_parts:
        cp = doc.add_paragraph(" | ".join(contact_parts))
        for r in cp.runs:
            r.font.size = Pt(9.5)

    doc.add_paragraph()

    doc.add_paragraph("Dear Hiring Manager,") if "Dear" not in body_text[:20] else None

    # Body paragraphs
    for para in body_text.split("\n\n"):
        para = para.strip()
        if not para:
            continue
        p = doc.add_paragraph(para)
        p.paragraph_format.space_after = Pt(8)

    doc.add_paragraph()
    closing = doc.add_paragraph("Sincerely,")
    closing.paragraph_format.space_after = Pt(2)
    name_close = doc.add_paragraph(resume_data.get("name", "").strip() or "Your Name")
    for r in name_close.runs:
        r.bold = True

    doc.save(output_path)
    return output_path
