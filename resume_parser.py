import os
import subprocess


def extract_text_from_file(filepath):
    """Extract plain text from a .pdf, .docx, or .txt resume file."""
    ext = os.path.splitext(filepath)[1].lower()

    if ext == ".txt":
        with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()

    if ext == ".docx":
        return _extract_docx(filepath)

    if ext == ".pdf":
        return _extract_pdf(filepath)

    raise ValueError(f"Unsupported file type: {ext}")


def _extract_docx(filepath):
    try:
        # Use pandoc if available for clean text extraction
        result = subprocess.run(
            ["pandoc", filepath, "-t", "plain"],
            capture_output=True, text=True, timeout=60
        )
        if result.returncode == 0 and result.stdout.strip():
            return result.stdout
    except (FileNotFoundError, subprocess.TimeoutExpired):
        pass

    # Fallback: python-docx
    try:
        from docx import Document
        doc = Document(filepath)
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text)
        return "\n".join(parts)
    except ImportError:
        raise RuntimeError("python-docx not installed and pandoc unavailable")


def _extract_pdf(filepath):
    try:
        result = subprocess.run(
            ["pdftotext", "-layout", filepath, "-"],
            capture_output=True, text=True, timeout=60
        )
        if result.returncode == 0 and result.stdout.strip():
            return result.stdout
    except (FileNotFoundError, subprocess.TimeoutExpired):
        pass

    # Fallback: pypdf
    try:
        from pypdf import PdfReader
        reader = PdfReader(filepath)
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    except ImportError:
        raise RuntimeError("pdftotext and pypdf unavailable - install poppler-utils or pypdf")
